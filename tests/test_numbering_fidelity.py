# SPDX-License-Identifier: MIT
"""Cluster D3 - LIST / NUMBERING definition fidelity (DOCX-ONLY).

Focused coverage for the numbering axis the complex-fixture suite does not reach:

- the structural per-level reader (``structure.num_per_level_facts`` /
  ``_abstract_num_per_level_facts`` / ``clone_abstract_num``) on a built shell;
- the fail-closed ``check_numbering_targets`` matrix (shell-defined PASS; undefined
  num_id / abstract_num_id ERROR; malformed numFmt ERROR; synthesized lvlText / indent
  ERROR; out-of-range indent ERROR; missing-numbering-part graceful fail-closed);
- multi-level capture (levels 0-2 declared) and the OOXML-inheritance non-error for the
  undeclared deeper levels.

Every shell is a tiny SYNTHETIC docx built in-test (no proprietary template). The engine
NEVER synthesizes a numbering definition - it can only REFERENCE the shell's numbering by
id and CLONE the shell's own w:abstractNum; this module proves the check enforces that.
"""

from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document  # noqa: E402

from brandkit.formats.docx import structure as docx_structure  # noqa: E402
from brandkit.formats.docx.structure import w  # noqa: E402
from brandkit.profile import schema  # noqa: E402
from brandkit.qa import checks_deterministic as cd  # noqa: E402


def _build_numbering_shell(path: Path, levels) -> Path:
    """A docx whose ``word/numbering.xml`` declares abstractNum 0 (``levels``) bound to
    w:num 1. ``levels`` is a list of ``(numFmt, lvlText, {indent attrs})`` per ilvl."""
    doc = Document()
    root = doc.part.numbering_part.element
    for child in list(root):
        root.remove(child)
    an = root.makeelement(w("abstractNum"), {w("abstractNumId"): "0"})
    root.append(an)
    for ilvl, (numfmt, lvltext, indent) in enumerate(levels):
        lvl = an.makeelement(w("lvl"), {w("ilvl"): str(ilvl)})
        an.append(lvl)
        nf = lvl.makeelement(w("numFmt"), {w("val"): numfmt})
        lvl.append(nf)
        lt = lvl.makeelement(w("lvlText"), {w("val"): lvltext})
        lvl.append(lt)
        if indent:
            ppr = lvl.makeelement(w("pPr"), {})
            lvl.append(ppr)
            ind = ppr.makeelement(w("ind"), {w(k): str(v) for k, v in indent.items()})
            ppr.append(ind)
    num = root.makeelement(w("num"), {w("numId"): "1"})
    root.append(num)
    aid = num.makeelement(w("abstractNumId"), {w("val"): "0"})
    num.append(aid)
    doc.save(path)
    return path


def _prof(numbering: dict) -> dict:
    """A minimal docx profile carrying ``numbering`` under one list role's appearance."""
    return {
        "kind": "docx",
        "roles": {
            "_index": ["list.bullet.1"],
            "list.bullet.1": {"appearance": {"numbering": numbering}},
        },
    }


class NumberingStructureReaderTest(unittest.TestCase):
    """The structural per-level readers surface the shell's OWN declared facts."""

    def test_per_level_facts_read_verbatim(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _build_numbering_shell(
                Path(td) / "s.docx",
                [("bullet", "", {"left": 720, "hanging": 360})],
            )
            doc = Document(shell)
            facts = docx_structure.num_per_level_facts(doc, "1")
            self.assertEqual(facts["num_id"], "1")
            self.assertEqual(facts["abstract_num_id"], "0")
            lvl0 = facts["per_level_facts"][0]
            self.assertEqual(lvl0["numFmt"], "bullet")
            self.assertEqual(lvl0["lvlText"], "")  # kept byte-for-byte
            self.assertEqual(lvl0["indent"], {"left": 720, "hanging": 360})

    def test_unknown_num_id_yields_none(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _build_numbering_shell(
                Path(td) / "s.docx", [("decimal", "%1.", {"left": 720})]
            )
            doc = Document(shell)
            self.assertIsNone(docx_structure.num_per_level_facts(doc, "9999"))

    def test_clone_abstract_num_is_verbatim_deep_copy(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _build_numbering_shell(
                Path(td) / "s.docx", [("decimal", "%1.", {"left": 720})]
            )
            doc = Document(shell)
            root = docx_structure._numbering_root(doc)
            clone = docx_structure.clone_abstract_num(root, "0")
            self.assertIsNotNone(clone)
            self.assertEqual(clone.get(w("abstractNumId")), "0")
            # A deep copy: mutating the clone never touches the live element.
            clone.set(w("abstractNumId"), "99")
            self.assertEqual(root.find(w("abstractNum")).get(w("abstractNumId")), "0")
            self.assertIsNone(docx_structure.clone_abstract_num(root, "77"))

    def test_multi_level_capture_only_declared_levels(self):
        # Levels 0-2 declared; the captured per_level_facts holds exactly those.
        with tempfile.TemporaryDirectory() as td:
            shell = _build_numbering_shell(
                Path(td) / "s.docx",
                [
                    ("decimal", "%1.", {"left": 720}),
                    ("lowerLetter", "%2.", {"left": 1440}),
                    ("lowerRoman", "%3.", {"left": 2160}),
                ],
            )
            doc = Document(shell)
            facts = docx_structure.num_per_level_facts(doc, "1")
            self.assertEqual(set(facts["per_level_facts"].keys()), {0, 1, 2})
            self.assertEqual(facts["per_level_facts"][2]["numFmt"], "lowerRoman")


class NumberingCheckMatrixTest(unittest.TestCase):
    """The fail-closed ``check_numbering_targets`` matrix."""

    def setUp(self):
        self._td = tempfile.TemporaryDirectory()
        self.addCleanup(self._td.cleanup)
        self.shell = _build_numbering_shell(
            Path(self._td.name) / "shell.docx",
            [("bullet", "", {"left": 720, "hanging": 360})],
        )
        self.valid = {
            "num_id": "1",
            "abstract_num_id": "0",
            "per_level_facts": {
                "0": {
                    "numFmt": "bullet",
                    "lvlText": "",
                    "indent": {"left": 720, "hanging": 360},
                }
            },
        }

    def _errors(self, numbering):
        return [
            f
            for f in cd.check_numbering_targets(self.shell, _prof(numbering))
            if f.severity == schema.Severity.ERROR.value
        ]

    def test_valid_shell_backed_numbering_passes(self):
        self.assertEqual(self._errors(self.valid), [])

    def test_non_docx_and_no_capture_are_noops(self):
        self.assertEqual(
            cd.check_numbering_targets(self.shell, {"kind": "pptx", "roles": {}}), []
        )
        self.assertEqual(
            cd.check_numbering_targets(
                self.shell, {"kind": "docx", "roles": {"_index": []}}
            ),
            [],
        )

    def test_undefined_num_id_errors(self):
        bad = dict(self.valid, num_id="9999")
        self.assertTrue(self._errors(bad))

    def test_undefined_abstract_num_id_errors(self):
        bad = dict(self.valid, abstract_num_id="77")
        self.assertTrue(self._errors(bad))

    def test_invalid_numfmt_errors(self):
        bad = {
            "num_id": "1",
            "abstract_num_id": "0",
            "per_level_facts": {"0": {"numFmt": "foobar"}},
        }
        self.assertTrue(self._errors(bad))

    def test_synthesized_lvltext_errors(self):
        # A lvlText that does NOT byte-match the shell's own for that level is rejected.
        bad = {
            "num_id": "1",
            "abstract_num_id": "0",
            "per_level_facts": {"0": {"lvlText": "%1)"}},
        }
        self.assertTrue(self._errors(bad))

    def test_out_of_range_indent_errors(self):
        bad = {
            "num_id": "1",
            "abstract_num_id": "0",
            "per_level_facts": {"0": {"indent": {"left": 999999}}},
        }
        self.assertTrue(self._errors(bad))

    def test_synthesized_indent_errors(self):
        # An in-range indent the shell never declared on that level is rejected.
        bad = {
            "num_id": "1",
            "abstract_num_id": "0",
            "per_level_facts": {"0": {"indent": {"left": 733}}},
        }
        self.assertTrue(self._errors(bad))

    def test_missing_numbering_part_fails_closed(self):
        # A shell with NO numbering part has empty inventories, so a referenced id can
        # never be a member -> the applied numbering is rejected (fail-closed).
        with tempfile.TemporaryDirectory() as td:
            empty = Path(td) / "empty.docx"
            doc = Document()
            root = doc.part.numbering_part.element
            for child in list(root):
                root.remove(child)
            doc.save(empty)
            self.assertTrue(
                [
                    f
                    for f in cd.check_numbering_targets(empty, _prof(self.valid))
                    if f.severity == schema.Severity.ERROR.value
                ]
            )


if __name__ == "__main__":
    unittest.main()
