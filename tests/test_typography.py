# SPDX-License-Identifier: MIT
"""Regression tests for brand typography capture (font family).

Covers the four layers of the feature:
  - capture: the dominant direct run font is recorded into theme.fonts.body and
    per-role appearance, and a no-dominant document captures nothing;
  - resolver: a role's own captured font wins; otherwise the document body font
    fills in (including for a missing-role stub);
  - apply: generated runs get the captured font as direct formatting, and a
    profile with NO captured typography leaves runs unfonted (no regression);
  - verify: a font the shell does not carry is an ERROR, a shell font is accepted,
    and an empty-appearance profile produces no finding.
"""

from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document

from brandkit.formats.docx import generate as docx_generate
from brandkit.formats.docx import typography
from brandkit.ir import model as ir
from brandkit.profile import schema
from brandkit.profile.resolver import ProfileResolver
from brandkit.qa import checks_deterministic


def _profile(theme=None, roles=None):
    prof = schema.build_envelope("docx", {"name": "t"})
    prof["surface"] = {"docx": {}}
    prof["roles"] = roles or {"_index": []}
    if theme is not None:
        prof["theme"] = theme
    return prof


def _shell(tmp_path, *, heading=True):
    shell = tmp_path / "shell.docx"
    d = Document()
    if heading:
        d.add_paragraph("x", style="Heading 1")
    d.save(shell)
    return shell


# ---------------------------------------------------------------------------
# capture
# ---------------------------------------------------------------------------
class CaptureTest(unittest.TestCase):
    def test_dominant_body_font_captured_into_theme(self):
        doc = Document()
        for _ in range(5):
            run = doc.add_paragraph().add_run("body text")
            run.font.name = "Roboto"
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertEqual(theme["fonts"]["body"]["latin"], "Roboto")
        self.assertGreaterEqual(theme["fonts"]["body"]["confidence"], 0.6)

    def test_per_role_font_captured_from_role_style(self):
        doc = Document()
        for _ in range(4):
            run = doc.add_paragraph(style="Heading 1").add_run("H")
            run.font.name = "Montserrat Black"
        # more body runs in a different font: the heading capture must stay
        # style-scoped (Montserrat) while the document body font is Roboto.
        for _ in range(8):
            doc.add_paragraph().add_run("b").font.name = "Roboto"
        roles = {
            "_index": ["heading.1"],
            "heading.1": {
                "resolver": {
                    "type": "named_style",
                    "style_id": "Heading1",
                    "style_name": "Heading 1",
                },
                "appearance": {},
            },
        }
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, roles, theme)
        self.assertEqual(
            roles["heading.1"]["appearance"]["font"]["latin"], "Montserrat Black"
        )
        self.assertEqual(theme["fonts"]["body"]["latin"], "Roboto")

    def test_no_explicit_font_captures_nothing(self):
        doc = Document()
        for _ in range(5):
            doc.add_paragraph().add_run("inherits the style font")  # no run font set
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertNotIn("body", theme["fonts"])

    def test_below_dominance_threshold_captures_nothing(self):
        doc = Document()
        fonts = ["Roboto", "Arial", "Times New Roman", "Courier"]  # 1/4 each, no winner
        for f in fonts:
            doc.add_paragraph().add_run("x").font.name = f
        theme = {"colors": {}, "fonts": {}}
        typography.capture_fonts(doc, {"_index": []}, theme)
        self.assertNotIn("body", theme["fonts"])


# ---------------------------------------------------------------------------
# resolver: role-specific font wins; body font is the fallback (incl. stub)
# ---------------------------------------------------------------------------
class ResolverAppearanceTest(unittest.TestCase):
    def _prof(self):
        return {
            "kind": "docx",
            "theme": {"colors": {}, "fonts": {"body": {"latin": "Roboto"}}},
            "roles": {
                "_index": ["heading.1"],
                "heading.1": {
                    "resolver": {"type": "named_style", "style_id": "Heading1"},
                    "appearance": {"font": {"latin": "Montserrat Black"}},
                    "status": "robust",
                    "confidence": 1.0,
                },
            },
        }

    def test_role_font_wins_over_body(self):
        op = ProfileResolver(self._prof()).resolve_role("heading.1")
        self.assertEqual(op.appearance["font"]["latin"], "Montserrat Black")

    def test_missing_role_stub_gets_body_font(self):
        op = ProfileResolver(self._prof()).resolve_role(
            "paragraph", fallback="paragraph"
        )
        self.assertEqual(op.appearance["font"]["latin"], "Roboto")

    def test_no_body_font_yields_empty_appearance(self):
        prof = self._prof()
        prof["theme"]["fonts"] = {}
        prof["roles"]["heading.1"]["appearance"] = {}
        op = ProfileResolver(prof).resolve_role("paragraph", fallback="paragraph")
        self.assertEqual(op.appearance, {})


# ---------------------------------------------------------------------------
# apply at generate time
# ---------------------------------------------------------------------------
class ApplyTest(unittest.TestCase):
    def test_body_font_applied_to_generated_runs(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={"colors": {}, "fonts": {"body": {"latin": "Roboto"}}}
            )
            idoc = ir.IntermediateDocument(
                blocks=[ir.Paragraph(runs=[{"t": "hello world"}])]
            )
            docx_generate.generate(prof, shell, idoc, out)
            fonts = {r.font.name for p in Document(out).paragraphs for r in p.runs}
            self.assertIn("Roboto", fonts)

    def test_role_font_wins_at_apply_time(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(
                theme={"colors": {}, "fonts": {"body": {"latin": "Roboto"}}},
                roles={
                    "_index": ["heading.1"],
                    "heading.1": {
                        "resolver": {
                            "type": "named_style",
                            "style_id": "Heading1",
                            "style_name": "Heading 1",
                        },
                        "appearance": {"font": {"latin": "Montserrat Black"}},
                    },
                },
            )
            idoc = ir.IntermediateDocument(
                blocks=[ir.Heading(level=1, runs=[{"t": "Title"}])]
            )
            docx_generate.generate(prof, shell, idoc, out)
            fonts = {r.font.name for p in Document(out).paragraphs for r in p.runs}
            self.assertIn("Montserrat Black", fonts)

    def test_no_captured_typography_leaves_runs_unfonted(self):
        # Regression: a profile with no theme.fonts.body and no role appearance must
        # produce runs with no direct font (exactly the pre-feature behavior).
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            out = Path(td) / "out.docx"
            prof = _profile(theme={"colors": {}, "fonts": {}})
            idoc = ir.IntermediateDocument(blocks=[ir.Paragraph(runs=[{"t": "plain"}])])
            docx_generate.generate(prof, shell, idoc, out)
            fonts = {r.font.name for p in Document(out).paragraphs for r in p.runs}
            self.assertEqual(fonts, {None})


# ---------------------------------------------------------------------------
# verify: fail-closed against the shell's available fonts
# ---------------------------------------------------------------------------
class AppearanceTargetsCheckTest(unittest.TestCase):
    def test_font_absent_from_shell_is_error(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            prof = _profile(
                theme={"colors": {}, "fonts": {"body": {"latin": "ZZZ Bogus Font"}}}
            )
            findings = checks_deterministic.check_appearance_targets(shell, prof)
            errs = [f for f in findings if f.check == "appearance_targets_exist"]
            self.assertTrue(errs)
            self.assertTrue(
                all(f.severity == schema.Severity.ERROR.value for f in errs)
            )

    def test_shell_font_is_accepted(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            # Arial is always in the docDefaults baseline / fontTable.
            prof = _profile(theme={"colors": {}, "fonts": {"body": {"latin": "Arial"}}})
            self.assertEqual(
                checks_deterministic.check_appearance_targets(shell, prof), []
            )

    def test_empty_appearance_profile_has_no_finding(self):
        with tempfile.TemporaryDirectory() as td:
            shell = _shell(Path(td))
            prof = _profile(theme={"colors": {}, "fonts": {}})
            self.assertEqual(
                checks_deterministic.check_appearance_targets(shell, prof), []
            )


if __name__ == "__main__":
    unittest.main()
